#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Writes a word document which contains information about the
configuration of the model, but also analyses outputs and performance.
This is currently mostly hardwired to the "var_kmax" project, but could
easily be tweaked.

This file is part of the TractLSM project.

Copyright (c) 2019 Manon E. B. Sabot

Please refer to the terms of the MIT License, which you should have
received along with the TractLSM.

"""

__title__ = "summary document of the model's configuration and outputs"
__author__ = "Manon E. B. Sabot"
__version__ = "1.0 (10.02.2018)"
__email__ = "m.e.b.sabot@gmail.com"

#=======================================================================

import warnings  # ignore these warnings
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=RuntimeWarning)

# import general modules
import argparse  # read in the user input
import os  # check for files, paths
import sys  # check for files, paths
import numpy as np  # array manipulations, math operators
import numpy.ma as ma  # masked arrays
import pandas as pd  # read/write dataframes, csv files
import xarray as xr  # read netcdf
from scipy.stats import percentileofscore as psc  # compute p-ranks
from geopy.geocoders import Nominatim  # match coordinates with country
import docx  # write out a word document
from docx import Document  # write out a word document

# first make sure that modules can be loaded from TractLSM
script_dir = os.path.dirname(os.path.realpath(sys.argv[0]))
sys.path.append(os.path.abspath(os.path.join(script_dir, '..')))

# own modules
from TractLSM import conv  # unit converter
from TractLSM import dparams as dp  # model parameters
from TractLSM.Utils import get_main_dir, read_csv  # locate & read


#=======================================================================

def main(projects):

    """
    Main function: writes a word document with information on:
                   (i) parameters and site-level input data used for the runs
                   (ii) general analysis about the output generated

    Arguments:
    ----------
    projects: array
        names of the repository which contain the output to analyse

    Returns:
    --------
    'summary_predicting_kmax.docx' in the output directory

    """

    write_word_doc(os.path.join('output', 'summary_predicting_kmax.docx'),
                   projects)

    return


#=======================================================================

# ~~~ Other functions are defined here ~~~

def reverse_geocode(lat, lon):

    """
    Finds the country associated with specific geographical coordinates.

    """

    location = Nominatim(user_agent="TractLSM").reverse('%s, %s' % (lat, lon))

    return location.address.rsplit(' ', 1)[1]


def kmax_values(sites, project, cLAI=None):

    """
    Retrieves the site values of kmax present within a specific project.

    Arguments:
    ----------
    sites: array
        names of the sites

    project: string
        name of the repository which contains the kmax data

    cLAI: array
        composite LAI to express kmax per area rather than per LAI

    Returns:
    --------
    kmax: array
        maximum hydraulic conductance [mmol m-2 s-1 MPa-1]

    """

    basedir = get_main_dir()

    while 'src' in basedir:
        basedir = os.path.dirname(get_main_dir())

    iname = os.path.join(os.path.join(os.path.join(basedir, 'input'),
                                      'projects'), project)

    kmax = np.zeros((len(sites), 3))
    confs = ['H', '', 'L']

    for i in range(3):

        for j in range(len(sites)):

            for file in os.listdir(iname):

                    if file.endswith('.csv') and ('actual' in file) and \
                       ('%s%s_' % (sites[j], confs[i]) in file):
                        df, __ = read_csv(os.path.join(iname, file))

                        if cLAI is None:
                            kmax[j, i] = df['kmax'].values[0]

                        else:
                            kmax[j, i] = df['kmax'].values[0] * cLAI[j]

    return kmax


def get_scores(df, sites, projects, kmax=None, climate=None, years=None):

    """
    Retrieves each of the yearly metrics and quantile ranks for all the
    sites and fluxes and projects.

    Arguments:
    ----------
    df: pandas dataframe
        dataframe containing all the performance scores

    sites: array
        site names

    projects: array
        names of the repository which contain the kmax data

    kmax: array
        kmax numbers for the calibration

    climate: array
        which is the best climate for each site

    years: array
        years to select the values from

    Returns:
    --------
    NMSE: array
        normalised mean square errors for each of the variables

    MAE: array
        mean absolute error

    SD: array
        standard deviation

    P5: array
        5th percentile

    P95: array
        95th percentile

    ranks: array
        quantile ranks

    """

    confs = ['', 'H', 'L']  # in opt-high-low order to calculate ranks
    projects = projects + ['control']

    # empty arrays to store the performance scores in
    NMSE = np.zeros((len(sites), len(projects), len(confs), 2, 2, 2))
    MAE = np.zeros((len(sites), len(projects), len(confs), 2, 2, 2))
    SD = np.zeros((len(sites), len(projects), len(confs), 2, 2, 2))
    P5 = np.zeros((len(sites), len(projects), len(confs), 2, 2, 2))
    P95 = np.zeros((len(sites), len(projects), len(confs), 2, 2, 2))

    # loop over ET and An
    for i in range(2):

        # loop over configurations
        for j in range(len(confs)):

            # loop over sites
            for k in range(len(sites)):

                for l in range(len(projects)):

                    subset = df[df['fname'].str.contains('%s%s_' % (sites[k],
                                                                    confs[j]))]

                    if 'sample' in projects[l]:
                        if kmax is not None:
                            subset = (df[df['fname'].str
                                      .contains('%s%s_' % (sites[k],
                                                           kmax[k]))])
                    if 'adjust' in projects[l]:
                        if climate is not None:
                            check = (subset['project']
                                     .str.contains(climate[k].lower()))

                            if any(check):
                                subset = subset[check]

                    sub = subset[subset['project'].str.contains(projects[l])]

                    for m in range(len(sub)):

                        ssub = sub.iloc[m]

                        if years is not None:
                            if any(str(yr) in ssub['fname'] for yr in years):
                                n = 0  # drought or not?

                            else:
                                n = 1

                        else:
                            n = 0

                        if i == 0:  # An
                            if np.isclose(NMSE[k, l, j, i, n, 0], 0.):
                                o = 0  # years

                            else:
                                o = 1

                            NMSE[k, l, j, i, n, o] = ssub['NMSE(An)']
                            MAE[k, l, j, i, n, o] = ssub['MAE(An)']
                            SD[k, l, j, i, n, o] = ssub['SD(An)']
                            P5[k, l, j, i, n, o] = ssub['P5(An)']
                            P95[k, l, j, i, n, o] = ssub['P95(An)']

                        else:  # ET
                            if np.isclose(NMSE[k, l, j, i, n, 0], 0.):
                                o = 0

                            else:
                                o = 1

                            NMSE[k, l, j, i, n, o] = ssub['NMSE(ET)']
                            MAE[k, l, j, i, n, o] = ssub['MAE(ET)']
                            SD[k, l, j, i, n, o] = ssub['SD(ET)']
                            P5[k, l, j, i, n, o] = ssub['P5(ET)']
                            P95[k, l, j, i, n, o] = ssub['P95(ET)']

    # deal with configs that don't apply (e.g. 'high' & 'low' to sample)
    for metric in [NMSE, MAE, SD, P5, P95]:
        metric[:, projects.index('var_kmax_sample'), 1:, :, :, :] = 0.
        metric[:, projects.index('control'), 1:, :, :, :] = 0.

    # mask zeros in array (only one year, or the control)
    NMSE = ma.masked_where(NMSE == 0., NMSE)
    NMSE = ma.mean(NMSE, axis=5)  # average over years
    MAE = ma.masked_where(MAE == 0., MAE)
    MAE = ma.mean(MAE, axis=5)
    SD = ma.masked_where(SD == 0., SD)
    SD = ma.mean(SD, axis=5)
    P5 = ma.masked_where(P5 == 0., P5)
    P5 = ma.mean(P5, axis=5)
    P95 = ma.masked_where(P95 == 0., P95)
    P95 = ma.mean(P95, axis=5)

    # compute the performance ranks now
    ranks = np.zeros((len(sites), len(projects), len(confs), 2, 2, 5))

    # loop over ET and An
    for i in range(2):

        # loop over sites
        for k in range(len(sites)):

            # loop over event types, i.e. drought or not
            for n in range(2):

                # loop over the metrics of performance
                for m, metric in enumerate([NMSE, MAE, SD, P5, P95]):

                    # average the metric within an event category
                    a = metric[k, :, :, i, n]
                    b = a[~ma.getmask(metric[k, :, :, i, n])]

                    # ranks computed in weak quantile regression
                    rk = np.array([psc(b, c, 'weak') / 100. for c in b])

                    track = 0  # tracker for the computed ranks

                    # loop over experiment types
                    for l in range(len(projects)):

                        # loop over configurations
                        for j in range(len(confs)):

                            # there are ranks where metrics are not nan!
                            if a[l, j] == b[track]:
                                ranks[k, l, j, i, n, m] = rk[track]
                                track += 1  # update the tracker

                                if track == len(b):  # all is filled
                                    break

    # average ranking across metrics
    ranks = ma.masked_where(ranks <= 0., ranks)
    ranks = ma.mean(ranks, axis=5)

    return NMSE, MAE, SD, P5, P95, ranks


def fluxes(sites, project, years=None, which='obs'):

    """
    Retrieves each of the yearly carbon and water flux data and performs
    all the appropriate unit conversions.

    Arguments:
    ----------
    sites: array
        site names

    project: string
        name of the repository which contains the flux input

    years: array
        years to select the values from

    which: string
        if 'obs' will look in the input files, else looks in the outputs

    Returns:
    --------
    GPP: array
        Gross Primary Productivity [g C m-2]

    ET: array
        Evapotranspiration [mm]

    ratio1: array
        instantaneous ratio of Es / ET [-]

    ratio2: array
        instantaneous ratio of E / ET [-]


    """

    if which == 'obs':
        basedir = get_main_dir()

        while 'src' in basedir:
            basedir = os.path.dirname(get_main_dir())

        iname = os.path.join(os.path.join(os.path.join(basedir, 'input'),
                                          'projects'), project)
        confs = ['']

        try:

            if len(years) > 1:
                GPP = np.zeros((len(sites, len(years))))
                ET = np.zeros((len(sites), len(years)))

        except TypeError:

            GPP = np.zeros(len(sites))
            ET = np.zeros(len(sites))

    else:
        basedir = get_main_dir()

        while 'src' in basedir:
            basedir = os.path.dirname(get_main_dir())

        iname = os.path.join(os.path.join(os.path.join(basedir, 'output'),
                                          'projects'), project)

        if which == 'calib':
            confs = ['']

            try:

                if len(years) > 1:
                    GPP = np.zeros((len(sites, len(years))))
                    ET = np.zeros((len(sites), len(years)))
                    ratio1 = np.zeros((len(sites), len(years)))
                    ratio2 = np.zeros((len(sites), len(years)))

            except TypeError:

                GPP = np.zeros(len(sites))
                ET = np.zeros(len(sites))
                ratio1 = np.zeros(len(sites))
                ratio2 = np.zeros(len(sites))

        else:
            confs = ['H', '', 'L']

            try:

                if len(years) > 1:
                    GPP = np.zeros((len(sites), 3, len(years)))
                    ET = np.zeros((len(sites), 3, len(years)))
                    ratio1 = np.zeros((len(sites), 3, len(years)))
                    ratio2 = np.zeros((len(sites), 3, len(years)))

            except TypeError:

                GPP = np.zeros((len(sites), 3))
                ET = np.zeros((len(sites), 3))
                ratio1 = np.zeros((len(sites), 3))
                ratio2 = np.zeros((len(sites), 3))

    for i in range(len(confs)):

        for j in range(len(sites)):

            for file in os.listdir(iname):

                find = True

                if file.endswith('.csv') and ('actual' in file) and \
                   ('%s%s_' % (sites[j], confs[i]) in file):
                    df, __ = read_csv(os.path.join(iname, file))

                    if years is not None:
                        try:

                            if len(years) > 1:
                                if any([str(int(y)) in file for y in years]):
                                    find = True

                                else:
                                    find = False

                        except TypeError:

                            if years not in file:
                                find = False

                    if find:
                        year = int(file.split('_')[-1].split('.csv')[0])

                        # restrict to subsets of dates
                        start = pd.to_datetime(year * 1000 + 1, format='%Y%j')
                        end = pd.to_datetime(year * 1000 + int(len(df) / 48.)
                                             + 1, format='%Y%j')

                        dates = pd.date_range(start=start, end=end, freq='D')
                        __, iunique = np.unique(dates.month, return_index=True)
                        unique_months = [dates.month[k] for k in
                                         sorted(iunique)]

                        # restrict between April - Nov for comparison
                        dates_2_keep = np.in1d(dates.month,
                                               unique_months[3:-2])
                        dates_2_keep = np.repeat(dates_2_keep, 48)[:len(df)]
                        df = df.iloc[dates_2_keep]

                        if which == 'obs':
                            mask = np.logical_and(df['PPFD'] > 50.,
                                                  df['LAI'] > 0.001)
                            GPP[j] += (df['GPP'][mask].sum() *
                                       conv.umolCpm2ps_2_gCpm2phlfhr)
                            ET[j] += (df['Qle'][mask].sum() *
                                      conv.Wpm2_2_mmphlfhr)

                        if which == 'sim':
                            GPP[j, i] += (df[df.filter(like='A(').columns]
                                          .sum() *
                                          conv.umolCpm2ps_2_gCpm2phlfhr)
                            ET_std = df['E(std)'] + df['Es(std)']
                            ET_psi = df['E(psi2)'] + df['Es(psi2)']
                            ET[j, i] += [(ET_std.sum() *
                                          conv.mmolH2Opm2ps_2_mmphlfhr),
                                         (ET_psi.sum() *
                                          conv.mmolH2Opm2ps_2_mmphlfhr)]
                            ratio1[j, i] += [(df['Es(std)'] / ET_std).mean(),
                                             (df['Es(psi2)'] / ET_psi).mean()]
                            ratio2[j, i] += [(df['E(std)'] / ET_std).mean(),
                                             (df['E(psi2)'] / ET_psi).mean()]

                        if which == 'calib':
                            GPP[j] += (df['A(psi2)'].sum()
                                       * conv.umolCpm2ps_2_gCpm2phlfhr)
                            ET_psi = df['E(psi2)'] + df['Es(psi2)']
                            ET[j] += (ET_psi.sum() *
                                      conv.mmolH2Opm2ps_2_mmphlfhr)
                            ratio1[j] += (df['Es(psi2)'] / ET_psi).mean()
                            ratio2[j] += (df['E(psi2)'] / ET_psi).mean()

    if which == 'obs':

        return GPP, ET

    else:

        return GPP, ET, ratio1, ratio2


def write_word_doc(fname, projects):

    """
    Writes the final word document. The order of the tables is as
    follows:
        1.  general PFT model parameters
        2.  photosynthesis model parameters

        3.  general site info
        4.  site climate and composite LAI
        5.  site soil parameters
        6.  site plant traits

        7.  site kmax climate values
        8.  site kmax climate range analysis
        9.  site kmax extreme:average climate ratio analysis

        10. ProfitMax NMSEs (close to zero is good)
        11. ProfitMax MAEs (close to zero is good)
        12. ProfitMax SD metrics (close to zero is good)
        13. ProfitMax P5 metrics (close to zero is good)
        14. ProfitMax P95 metrics (close to zero is good)
        15. % improvement on the Calibration metrics compared to the
            Control

        16. ProfitMax "optimal" behaviour quantile ranks
        17. ProfitMax "high" behaviour quantile ranks
        18. ProfitMax "low" behaviour quantile ranks

        19. Total GPP between April-Nov
        20. % improvement on the total GPP compared to the Control
        21. Total ET between April-Nov
        22. % improvement on the total ET compared to the Control
        23. average Es / ET and average E / ET ratios between April-Nov

    Arguments:
    ----------
    fname: string
        name of the summary file to save

    projects: array
        names of the repository which contain the kmax data

    Returns:
    --------
    fname in the directory which harbours TractLSM

    """

    # useful directories
    basedir = get_main_dir()

    while 'src' in basedir:
        basedir = os.path.dirname(get_main_dir())

    Rsites = os.path.join(os.path.join(basedir, 'input'), 'fluxsites')
    Rmet = os.path.join(Rsites, 'met')
    Rout = os.path.join(os.path.join(basedir, 'output'), 'projects')

    # site info
    df, __ = read_csv(os.path.join(Rsites, 'info_sites.csv'))
    df.set_index('Site', inplace=True)

    # declare document
    doc = Document()

    doc.add_heading("Summary of the TractLSM outputs", 0)

    description = """This document summarizes:\n
                     (i) parameters and site-level input data used for the runs
                     (ii) general analysis about the output generated
                  """
    p = doc.add_paragraph(description)
    p.add_run('\nThe following projects are accounted for:\n')

    for project in projects:

        p.add_run('     %s\n' % (project)).bold = True

        if 'sample' in project:
            try:
                fkmax = open(os.path.join(os.path.join(Rout,
                                                       'var_kmax_best_calib'),
                                          'best.txt'), 'r')
                flines = fkmax.readlines()
                kmax_calib = [None] * len(sites)

                for l in flines:

                    site = l.split(':')[0].strip()
                    k = l.split(':')[1].strip()
                    kmax_calib[sites.index(site)] = k

                if not any(kmax_calib):
                    kmax_calib = None

            except IOError:
                kmax_calib = None

        if 'adjust' in project:
            try:  # look for "best" climate
                fclimate = open(os.path.join(os.path.join(Rout,
                                'var_kmax_best_climate'), 'best.txt'), 'r')
                flines = fclimate.readlines()
                climate = [None] * len(sites)

                for l in flines:

                    site = l.split(':')[0].strip()
                    cl = l.split(':')[1].strip()

                    if 'average' in cl:
                        climate[sites.index(site)] = 'Average'

                    else:
                        climate[sites.index(site)] = 'Extreme'

                if not any(climate):
                    climate = None

            except IOError:
                climate = None

    doc.add_heading('General (PFT-based) Model parameters', level=1)
    PFTs = np.unique(df['PFT'][~pd.isna(df['PFT'])])

    t = doc.add_table(rows=9, cols=len(PFTs) + 3)
    t.cell(0, 0).text = 'Abbreviation'
    t.cell(0, 1).text = 'Parameter description'
    t.cell(0, 2).text = 'Value'
    t.cell(0, 2 + len(PFTs)).text = 'Unit'

    track = 0

    for PFT in PFTs:

        t.cell(1, 2 + track).text = str(PFT)

        track += 1

    # rendering of the headers
    t.cell(0, 0).merge(t.cell(1, 0))
    t.cell(0, 1).merge(t.cell(1, 1))
    t.cell(0, 2 + len(PFTs)).merge(t.cell(1, 2 + len(PFTs)))

    for i in range(3, 2 + len(PFTs)):

        t.cell(0, 2).merge(t.cell(0, i))
        t.cell(5, 2).merge(t.cell(5, i))
        t.cell(6, 2).merge(t.cell(6, i))
        t.cell(7, 2).merge(t.cell(7, i))

    # fill the table
    t.cell(2, 0).text = 'w_l_max'
    t.cell(2, 1).text = 'maximum leaf width'
    t.cell(2, 2 + len(PFTs)).text = 'm'

    for i in range(len(PFTs)):

        t.cell(2, 2 + i).text = '{0:.3f}'.format(np.unique(df['max_leaf_width']
                                                 [df['PFT'] == PFTs[i]])[0])

    t.cell(3, 0).text = 'can_int'
    t.cell(3, 1).text = 'canopy intercept'
    t.cell(3, 2 + len(PFTs)).text = '-'

    for i in range(len(PFTs)):

        t.cell(3, 2 + i).text = ('{0:.3f}'
                                 .format(np.unique(df['canopy_intercept']
                                         [df['PFT'] == PFTs[i]])[0]))

    t.cell(4, 0).text = 'alpha_l'
    t.cell(4, 1).text = 'leaf albedo'
    t.cell(4, 2 + len(PFTs)).text = '-'

    for i in range(len(PFTs)):

        t.cell(4, 2 + i).text = '{0:.3f}'.format(np.unique(df['albedo_l']
                                                 [df['PFT'] == PFTs[i]])[0])

    t.cell(5, 0).text = 'alpha_ws'
    t.cell(5, 1).text = 'wet soil albedo'
    t.cell(5, 2 + len(PFTs)).text = '-'
    t.cell(5, 2 + i).text = '{0:.3f}'.format(dp.albedo_ws)

    t.cell(6, 0).text = 'alpha_ds'
    t.cell(6, 1).text = 'dry soil albedo'
    t.cell(6, 2 + len(PFTs)).text = '-'
    t.cell(6, 2 + i).text = '{0:.3f}'.format(dp.albedo_ds)

    t.cell(7, 0).text = 'eps_l'
    t.cell(7, 1).text = 'leaf emissivity'
    t.cell(7, 2 + len(PFTs)).text = '-'
    t.cell(7, 2 + i).text = '{0:.3f}'.format(dp.eps_l)

    t.cell(8, 0).text = 'eps_s'
    t.cell(8, 1).text = 'soil emissivity'
    t.cell(8, 2 + len(PFTs)).text = '-'
    t.cell(8, 2 + i).text = '{0:.3f}'.format(dp.eps_s)

    doc.add_heading('Photosynthetis Model parameters', level=1)

    t = doc.add_table(rows=19, cols=4)
    t.cell(0, 0).text = 'Abbreviation'
    t.cell(0, 1).text = 'Parameter description'
    t.cell(0, 2).text = 'Value'
    t.cell(0, 3).text = 'Unit'

    # fill the table
    t.cell(1, 0).text = 'CO2'
    t.cell(1, 1).text = 'atmospheric CO2 concentration'
    t.cell(1, 2).text = '{0:.2f}'.format(dp.CO2)
    t.cell(1, 3).text = 'Pa'

    t.cell(2, 0).text = 'O2'
    t.cell(2, 1).text = 'atmospheric O2 concentration'
    t.cell(2, 2).text = '{0:.2f}'.format(dp.O2)
    t.cell(2, 3).text = 'kPa'

    t.cell(3, 0).text = 'gamma star'
    t.cell(3, 1).text = 'CO2 compensation point @ 25 degC'
    t.cell(3, 2).text = '{0:.2f}'.format(dp.gamstar25)
    t.cell(3, 3).text = 'Pa'

    t.cell(4, 0).text = 'Kc'
    t.cell(4, 1).text = 'Michaelis-Menton constant for carboxylation'
    t.cell(4, 2).text = '{0:.2f}'.format(dp.Kc25)
    t.cell(4, 3).text = 'Pa'

    t.cell(5, 0).text = 'Ko'
    t.cell(5, 1).text = 'Michaelis-Menton constant for oxygenation'
    t.cell(5, 2).text = '{0:.2f}'.format(dp.Ko25)
    t.cell(5, 3).text = 'kPa'

    t.cell(6, 0).text = 'J:V'
    t.cell(6, 1).text = 'Jmax25 to Vmax25 ratio'
    t.cell(6, 2).text = '{0:.2f}'.format(dp.JV)
    t.cell(6, 3).text = '-'

    t.cell(7, 0).text = 'alpha'
    t.cell(7, 1).text = 'quantum yield of electron transport'
    t.cell(7, 2).text = '{0:.2f}'.format(dp.alpha)
    t.cell(7, 3).text = 'mol(photon).mol(e-)-1'

    t.cell(8, 0).text = 'c1'
    t.cell(8, 1).text = 'curvature of the light response'
    t.cell(8, 2).text = '{0:.4f}'.format(dp.c1)
    t.cell(8, 3).text = '-'

    t.cell(9, 0).text = 'c2'
    t.cell(9, 1).text = 'transition Je vs Jc'
    t.cell(9, 2).text = '{0:.2f}'.format(dp.c2)
    t.cell(9, 3).text = '-'

    t.cell(10, 0).text = 'Ec'
    t.cell(10, 1).text = 'energy of activation of the carboxylation'
    t.cell(10, 2).text = '{0:.2f}'.format(dp.Ec)
    t.cell(10, 3).text = 'J.mol-1'

    t.cell(11, 0).text = 'Eo'
    t.cell(11, 1).text = 'energy of activation of the oxygenation'
    t.cell(11, 2).text = '{0:.2f}'.format(dp.Eo)
    t.cell(11, 3).text = 'J.mol-1'

    t.cell(12, 0).text = 'Ev'
    t.cell(12, 1).text = 'energy of activation of Vcmax'
    t.cell(12, 2).text = '{0:.2f}'.format(dp.Ev)
    t.cell(12, 3).text = 'J.mol-1'

    t.cell(13, 0).text = 'Ej'
    t.cell(13, 1).text = 'energy of activation of Jmax'
    t.cell(13, 2).text = '{0:.2f}'.format(dp.Ej)
    t.cell(13, 3).text = 'J.mol-1'

    t.cell(14, 0).text = 'Egamm'
    t.cell(14, 1).text = 'energy of activation of the CO2 compensation point'
    t.cell(14, 2).text = '{0:.2f}'.format(dp.Egamstar)
    t.cell(14, 3).text = 'J.mol-1'

    t.cell(15, 0).text = 'deltaSv'
    t.cell(15, 1).text = 'Vcmax entropy factor'
    t.cell(15, 2).text = '{0:.2f}'.format(dp.deltaSv)
    t.cell(15, 3).text = 'J.mol-1.K-1'

    t.cell(16, 0).text = 'deltaSj'
    t.cell(16, 1).text = 'Jmax entropy factor'
    t.cell(16, 2).text = '{0:.2f}'.format(dp.deltaSj)
    t.cell(16, 3).text = 'J.mol-1.K-1'

    t.cell(17, 0).text = 'Hdv'
    t.cell(17, 1).text = 'Vcmax rate of decrease above the optimum T'
    t.cell(17, 2).text = '{0:.2f}'.format(dp.Hdv)
    t.cell(17, 3).text = 'J.mol-1'

    t.cell(18, 0).text = 'Hdj'
    t.cell(18, 1).text = 'Jmax rate of decrease above the optimum T'
    t.cell(18, 2).text = '{0:.2f}'.format(dp.Hdj)
    t.cell(18, 3).text = 'J.mol-1'

    doc.add_heading('General Site Info', level=1)

    t = doc.add_table(rows=len(sites) + 1, cols=8)
    t.cell(0, 0).text = 'Site Name'
    t.cell(0, 1).text = 'Country'
    t.cell(0, 2).text = 'Latitude'
    t.cell(0, 3).text = 'Longitude'
    t.cell(0, 4).text = 'PFT'
    t.cell(0, 5).text = 'Dominant Species'
    t.cell(0, 6).text = 'Data'
    t.cell(0, 7).text = 'Reference'

    for i in range(len(sites)):

        site = sites[i]

        ds = xr.open_dataset(os.path.join(Rmet,
                                          '%sFluxnet_met.nc' % (site)))
        lat = ds.latitude.values[0][0]
        lon = ds.longitude.values[0][0]

        if ds.attrs['Fluxnet_dataset_version'] == 'n/a':
            dataset = 'LT'

        else:
            dataset = 'FN'

        Spp = df.loc[site, 'Dominant Species']

        if len(Spp.split()) > 2:
            Spp = Spp.rsplit(' ', 1)[0]

        t.cell(i + 1, 0).text = site
        try:
            t.cell(i + 1, 1).text = reverse_geocode(lat, lon)

        except Exception:
            pass

        t.cell(i + 1, 2).text = '{0:.2f}'.format(lat)
        t.cell(i + 1, 3).text = '{0:.2f}'.format(lon)
        t.cell(i + 1, 4).text = df.loc[site, 'PFT']
        t.cell(i + 1, 5).text = Spp
        t.cell(i + 1, 6).text = dataset
        t.cell(i + 1, 7).text = ''

    doc.add_heading('Site Climate and composite LAI', level=1)

    t = doc.add_table(rows=len(sites) + 1, cols=8)
    t.cell(0, 0).text = 'Site Name'
    t.cell(0, 1).text = 'Koeppen Climate Class'
    t.cell(0, 2).text = 'MAP\n(mm/y)'
    t.cell(0, 3).text = 'Tair,avg\n(degC)'
    t.cell(0, 4).text = 'Davg\n(kPa)'
    t.cell(0, 5).text = 'Tair,xx\n(degC)'
    t.cell(0, 6).text = 'Dxx\n(kPa)'
    t.cell(0, 7).text = 'Composite LAI\n(m2/m2)'

    cLAI = []  # store the composite LAI for kmax tables below

    for i in range(len(sites)):

        site = sites[i]

        t.cell(i + 1, 0).text = site
        t.cell(i + 1, 1).text = df.loc[site, 'Koeppen Climate']
        t.cell(i + 1, 2).text = '{0:.2f}'.format(df.loc[site, 'CRU MAP'])
        t.cell(i + 1, 3).text = '{0:.2f}'.format(df.loc[site, 'CRU cTair'])
        t.cell(i + 1, 4).text = '{0:.2f}'.format(df.loc[site, 'CRU cVPD'])
        t.cell(i + 1, 5).text = '{0:.2f}'.format(df.loc[site, 'CRU cTxx'])
        t.cell(i + 1, 6).text = '{0:.2f}'.format(df.loc[site, 'CRU cVPDxx'])
        t.cell(i + 1, 7).text = '{0:.2f}'.format(df.loc[site, 'Composite LAI'])

        cLAI += [df.loc[site, 'Composite LAI']]

    doc.add_heading('Site Soil Parameters', level=1)

    t = doc.add_table(rows=len(sites) + 1, cols=9)
    t.cell(0, 0).text = 'Site Name'
    t.cell(0, 1).text = 'Psie\n(-MPa)'
    t.cell(0, 2).text = 'k_s,sat\n(m s-1)'
    t.cell(0, 3).text = 'theta_sat\n(m3/m3)'
    t.cell(0, 4).text = 'theta_fc\n(m3/m3)'
    t.cell(0, 5).text = 'theta_wp\n(m3/m3)'
    t.cell(0, 6).text = 'b\n(-)'
    t.cell(0, 7).text = 'r_soil\n(-)'
    t.cell(0, 8).text = 'Z_total\n(m)'

    for i in range(len(sites)):

        site = sites[i]

        t.cell(i + 1, 0).text = site
        t.cell(i + 1, 1).text = '{0:.2e}'.format(-df.loc[site, 'Psie'])
        t.cell(i + 1, 2).text = '{0:.2e}'.format(-df.loc[site, 'hyds'])
        t.cell(i + 1, 3).text = '{0:.2e}'.format(df.loc[site, 'theta_sat'])
        t.cell(i + 1, 4).text = '{0:.2e}'.format(df.loc[site, 'fc'])
        t.cell(i + 1, 5).text = '{0:.2e}'.format(df.loc[site, 'pwp'])
        t.cell(i + 1, 6).text = '{0:.2f}'.format(df.loc[site, 'bch'])
        t.cell(i + 1, 7).text = '{0:.2f}'.format(df.loc[site, 'r_soil'])
        t.cell(i + 1, 8).text = '{0:.2f}'.format(df.loc[site, 'Zbottom'])

    doc.add_heading('Site Plant Traits', level=1)

    t = doc.add_table(rows=len(sites) + 1, cols=6)
    t.cell(0, 0).text = 'Site Name'
    t.cell(0, 1).text = 'Vcmax25\n(umol/m2/s)'
    t.cell(0, 2).text = 'Rd25\n(umol/m2/s)'
    t.cell(0, 3).text = 'g1\n(kPa1/2)'
    t.cell(0, 4).text = 'P50\n(-MPa)'
    t.cell(0, 5).text = 'P88\n(-MPa)'

    for i in range(len(sites)):

        site = sites[i]

        t.cell(i + 1, 0).text = site
        t.cell(i + 1, 1).text = '{0:.2f}'.format(df.loc[site, 'Vmax25'])
        t.cell(i + 1, 2).text = '{0:.2f}'.format(df.loc[site, 'Rlref'])
        t.cell(i + 1, 3).text = '{0:.2f}'.format(df.loc[site, 'g1'])
        t.cell(i + 1, 4).text = '{0:.2f}'.format(df.loc[site, 'P50'])
        t.cell(i + 1, 5).text = '{0:.2f}'.format(df.loc[site, 'P88'])

    doc.add_heading('Site kmax values', level=1)

    # retrieve the kmax values
    kmax_avg = None
    kmax_xx = None
    kmax = None

    for project in projects:

        if 'average' in project:
            kmax_avg = kmax_values(sites, project, cLAI=cLAI)

        if 'extreme' in project:
            kmax_xx = kmax_values(sites, project, cLAI=cLAI)

    t = doc.add_table(rows=len(sites) + 2, cols=7)
    t.cell(0, 0).text = 'Site Name'
    t.cell(0, 1).text = 'average scenario'
    t.cell(0, 4).text = 'extreme scenario'
    t.cell(1, 1).text = 'kmax,high\n(mmol/m2/s/MPa)'
    t.cell(1, 2).text = 'kmax,optimal\n(mmol/m2/s/MPa)'
    t.cell(1, 3).text = 'kmax,low\n(mmol/m2/s/MPa)'
    t.cell(1, 4).text = 'kmax,high\n(mmol/m2/s/MPa)'
    t.cell(1, 5).text = 'kmax,optimal\n(mmol/m2/s/MPa)'
    t.cell(1, 6).text = 'kmax,low\n(mmol/m2/s/MPa)'

    # rendering of the headers
    t.cell(0, 0).merge(t.cell(1, 0))
    t.cell(0, 1).merge(t.cell(0, 2))
    t.cell(0, 1).merge(t.cell(0, 3))
    t.cell(0, 4).merge(t.cell(0, 5))
    t.cell(0, 4).merge(t.cell(0, 6))

    for i in range(len(sites)):

        site = sites[i]
        t.cell(i + 2, 0).text = site

        for project in projects:

            if 'average' in project:
                loc = 1
                kmax = kmax_avg[i]

            if 'extreme' in project:
                loc = 4
                kmax = kmax_xx[i]

            if kmax is not None:

                for j in range(len(kmax)):

                    t.cell(i + 2, loc + j).text = '{0:.2f}'.format(kmax[j])

    doc.add_heading('Site kmax range analysis', level=1)

    t = doc.add_table(rows=len(sites) + 4, cols=7)
    t.cell(0, 0).text = 'Site Name'
    t.cell(0, 1).text = 'average scenario'
    t.cell(0, 3).text = 'extreme scenario'
    t.cell(0, 5).text = 'overall'
    t.cell(1, 1).text = 'min'
    t.cell(1, 2).text = 'max'
    t.cell(1, 3).text = 'min'
    t.cell(1, 4).text = 'max'
    t.cell(1, 5).text = 'min'
    t.cell(1, 6).text = 'max'

    # rendering of the headers
    t.cell(0, 0).merge(t.cell(1, 0))
    t.cell(0, 1).merge(t.cell(0, 2))
    t.cell(0, 3).merge(t.cell(0, 4))
    t.cell(0, 5).merge(t.cell(0, 6))

    for i in range(len(sites)):

        site = sites[i]
        t.cell(i + 2, 0).text = site

        if kmax_avg is not None:
            t.cell(i + 2, 1).text = '{0:.2f}'.format(np.amin(kmax_avg[i]))
            t.cell(i + 2, 2).text = '{0:.2f}'.format(np.amax(kmax_avg[i]))

        if kmax_xx is not None:
            t.cell(i + 2, 3).text = '{0:.2f}'.format(np.amin(kmax_xx[i]))
            t.cell(i + 2, 4).text = '{0:.2f}'.format(np.amax(kmax_xx[i]))

        if (kmax_avg is not None) and (kmax_xx is not None):
            kmax = np.concatenate((kmax_avg[i], kmax_xx[i]))

        elif (kmax_avg is not None) and (kmax_xx is None):
            kmax = kmax_avg[i]

        elif (kmax_xx is not None) and (kmax_avg is None):
            kmax = kmax_xx[i]

        if (kmax_avg is None) and (kmax_xx is None):
            t.cell(i + 2, 5).text = ''
            t.cell(i + 2, 6).text = ''

        else:
            t.cell(i + 2, 5).text = '{0:.2f}'.format(np.amin(kmax))
            t.cell(i + 2, 6).text = '{0:.2f}'.format(np.amax(kmax))

    t.cell(len(sites) + 3, 0).text = 'across sites'

    if kmax_avg is not None:
        t.cell(len(sites) + 3, 1).text = '{0:.2f}'.format(np.amin(kmax_avg))
        t.cell(len(sites) + 3, 2).text = '{0:.2f}'.format(np.amax(kmax_avg))

    if kmax_xx is not None:
        t.cell(len(sites) + 3, 3).text = '{0:.2f}'.format(np.amin(kmax_xx))
        t.cell(len(sites) + 3, 4).text = '{0:.2f}'.format(np.amax(kmax_xx))

    if (kmax_avg is not None) and (kmax_xx is not None):
        kmax = np.concatenate((kmax_avg, kmax_xx))
        t.cell(len(sites) + 3, 5).text = '{0:.2f}'.format(np.amin(kmax))
        t.cell(len(sites) + 3, 6).text = '{0:.2f}'.format(np.amax(kmax))

    doc.add_heading('Site kmax ratio analysis', level=1)

    t = doc.add_table(rows=len(sites) + 4, cols=5)
    t.cell(0, 0).text = 'Site Name'
    t.cell(0, 1).text = 'extreme:average'
    t.cell(1, 1).text = 'high'
    t.cell(1, 2).text = 'optimal'
    t.cell(1, 3).text = 'low'
    t.cell(1, 4).text = 'overall'

    # rendering of the headers
    t.cell(0, 0).merge(t.cell(1, 0))
    t.cell(0, 1).merge(t.cell(0, 2))
    t.cell(0, 1).merge(t.cell(0, 3))
    t.cell(0, 1).merge(t.cell(0, 4))

    if (kmax_avg is not None) and (kmax_xx is not None):
        ratios = kmax_xx / kmax_avg

    for i in range(len(sites)):

        site = sites[i]
        t.cell(i + 2, 0).text = site

        if (kmax_avg is not None) and (kmax_xx is not None):
            t.cell(i + 2, 1).text = '{0:.2f}'.format(ratios[i, 0])
            t.cell(i + 2, 2).text = '{0:.2f}'.format(ratios[i, 1])
            t.cell(i + 2, 3).text = '{0:.2f}'.format(ratios[i, 2])
            t.cell(i + 2, 4).text = '{0:.2f}'.format(np.mean(ratios[i]))

    t.cell(len(sites) + 3, 0).text = 'across sites'

    if (kmax_avg is not None) and (kmax_xx is not None):
        t.cell(len(sites) + 3, 1).text = ('{0:.2f}'
                                          .format(np.mean(ratios[:, 0])))
        t.cell(len(sites) + 3, 2).text = ('{0:.2f}'
                                          .format(np.mean(ratios[:, 1])))
        t.cell(len(sites) + 3, 3).text = ('{0:.2f}'
                                          .format(np.mean(ratios[:, 2])))
        t.cell(len(sites) + 3, 4).text = ('{0:.2f}'
                                          .format(np.mean(ratios[:, :])))

    # performance scores
    dfp, __ = read_csv(os.path.join(os.path.join(Rout, 'var_kmax'),
                       'perf_scores.csv'))
    dfp.set_index('Site', inplace=True)

    scores = (get_scores(dfp, sites, projects, kmax=kmax_calib,
                         climate=climate, years=[2003, 2006]))

    track = 0

    for perf in ['NMSE', 'MAE', 'SD', 'P5', 'P95']:

        doc.add_heading('ProfitMax optimal %ss' % (perf), level=1)
        score = scores[track]

        t = doc.add_table(rows=len(sites) * 2 + 7, cols=10)
        t.cell(0, 0).text = 'Site Name'
        t.cell(0, 1).text = 'Flux'
        t.cell(0, 2).text = '2003 & 2006'
        t.cell(0, 6).text = '2002 & 2005'
        t.cell(1, 2).text = 'average scenario'
        t.cell(1, 3).text = 'extreme scenario'
        t.cell(1, 4).text = 'control'
        t.cell(1, 5).text = 'calib'
        t.cell(1, 6).text = 'average scenario'
        t.cell(1, 7).text = 'extreme scenario'
        t.cell(1, 8).text = 'control'
        t.cell(1, 9).text = 'calib'

        # rendering of the headers
        t.cell(0, 0).merge(t.cell(1, 0))
        t.cell(0, 2).merge(t.cell(0, 3))
        t.cell(0, 3).merge(t.cell(0, 4))
        t.cell(0, 4).merge(t.cell(0, 5))
        t.cell(0, 6).merge(t.cell(0, 7))
        t.cell(0, 7).merge(t.cell(0, 8))
        t.cell(0, 8).merge(t.cell(0, 9))

        for i in range(len(sites)):

            site = sites[i]
            t.cell(i * 2 + 2, 0).text = site

            ds = dfp.loc[site]

            t.cell(i * 2 + 2, 1).text = 'GPP'
            t.cell(i * 2 + 3, 1).text = 'ET'

            for j in range(2):  # GPP, ET

                for k in range(2):  # events

                    if k == 0:
                        bloc = 2

                    else:
                        bloc = 6

                    t.cell(i * 2 + j + 2, bloc).text = \
                        '{0:.2f}'.format(score[i, 0, 0, j, k])
                    t.cell(i * 2 + j + 2, bloc + 1).text = \
                        '{0:.2f}'.format(score[i, 1, 0, j, k])
                    t.cell(i * 2 + j + 2, bloc + 2).text = \
                        '{0:.2f}'.format(score[i, len(score[0])-1, 0, j, k])

                    if kmax_calib is not None:
                        t.cell(i * 2 + j + 2, bloc + 3).text = \
                            '{0:.2f}'.format(score[i, 2, 0, j, k])

        t.cell(len(sites) * 2 + 2, 0).text = 'across sites'
        t.cell(len(sites) * 2 + 2, 1).text = 'GPP'
        t.cell(len(sites) * 2 + 3, 1).text = 'ET'

        for j in range(2):  # GPP, ET

            for k in range(2):  # events

                if k == 0:
                    bloc = 2

                else:
                    bloc = 6

                t.cell(len(sites) * 2 + j + 2, bloc).text = \
                    '{0:.2f}'.format(np.mean(score[:, 0, 0, j, k]))
                t.cell(len(sites) * 2 + j + 2, bloc + 1).text = \
                    '{0:.2f}'.format(np.mean(score[:, 1, 0, j, k]))
                t.cell(len(sites) * 2 + j + 2, bloc + 2).text = \
                    '{0:.2f}'.format(np.mean(score[:,
                                                   len(score[0])-1, 0, j, k]))

                if kmax_calib is not None:
                    t.cell(len(sites) * 2 + j + 2, bloc + 3).text = \
                        '{0:.2f}'.format(np.mean(score[:, 2, 0, j, k]))

        t.cell((len(sites) + 1) * 2 + 3, 0).text = 'overall'
        t.cell((len(sites) + 1) * 2 + 3, 1).text = 'GPP'
        t.cell((len(sites) + 1) * 2 + 4, 1).text = 'ET'

        for j in range(2):  # GPP, ET

            t.cell((len(sites) + 1) * 2 + j + 3, 2).text = \
                '{0:.2f}'.format(np.mean(score[:, 0, 0, j, :]))
            t.cell((len(sites) + 1) * 2 + j + 3, 3).text = \
                '{0:.2f}'.format(np.mean(score[:, 1, 0, j, :]))
            t.cell((len(sites) + 1) * 2 + j + 3, 4).text = \
                '{0:.2f}'.format(np.mean(score[:, len(score[0])-1, 0, j, :]))

            if kmax_calib is not None:
                t.cell((len(sites) + 1) * 2 + j + 3, 5).text = \
                    '{0:.2f}'.format(np.mean(score[:, 2, 0, j, :]))

        track += 1

    doc.add_heading('% overall improvement on Control metrics', level=1)

    t = doc.add_table(rows=11, cols=4)
    t.cell(0, 0).text = 'metric'
    t.cell(0, 1).text = 'flux'
    t.cell(0, 2).text = 'climate'
    t.cell(0, 3).text = 'calib'

    track1 = 0
    track2 = 1

    for perf in ['NMSE', 'MAE', 'SD', 'P5', 'P95']:

        score = ma.filled(scores[track1], 0.)  # 0. where masked values

        t.cell(track2, 0).text = perf
        t.cell(track2, 0).merge(t.cell(track2 + 1, 0))

        for j in range(2):  # GPP, ET

            if j == 0:
                t.cell(track2, 1).text = 'GPP'

            else:
                t.cell(track2, 1).text = 'ET'

            ctrl = np.mean(score[:, len(score[0])-1, 0, j, :])

            if (kmax_avg is not None) and (kmax_xx is not None):
                climate = np.mean(score[:, 0, 0, j, :] + score[:, 1, 0, j, :])
                climate = (abs(ctrl) - abs(climate)) / abs(ctrl)
                t.cell(track2, 2).text = '{0:.2f}'.format(100. * climate)

            if kmax_calib is not None:
                cal = np.mean(score[:, 2, 0, j, :])
                cal = (abs(ctrl) - abs(cal)) / abs(ctrl)
                t.cell(track2, 3).text = '{0:.2f}'.format(100. * cal)

            track2 += 1

        track1 += 1

    track = 0

    for conf in ['optimal', 'high', 'low']:

        doc.add_heading('ProfitMax %s Quantile Ranks' % (conf), level=1)
        score = ma.filled(scores[-1][:, :, track, :, :], 0.)  # masked

        if conf == 'optimal':
            t = doc.add_table(rows=len(sites) * 2 + 11, cols=8)

        else:
            t = doc.add_table(rows=len(sites) * 2 + 2, cols=8)

        t.cell(0, 0).text = 'Site Name'
        t.cell(0, 1).text = 'Flux'
        t.cell(0, 2).text = '2003 & 2006'
        t.cell(0, 6).text = '2002 & 2005'
        t.cell(1, 2).text = 'climate'
        t.cell(1, 3).text = 'control'
        t.cell(1, 4).text = 'calib'
        t.cell(1, 5).text = 'climate'
        t.cell(1, 6).text = 'control'
        t.cell(1, 7).text = 'calib'

        # rendering of the headers
        t.cell(0, 0).merge(t.cell(1, 0))
        t.cell(0, 2).merge(t.cell(0, 3))
        t.cell(0, 3).merge(t.cell(0, 4))
        t.cell(0, 4).merge(t.cell(0, 5))
        t.cell(0, 6).merge(t.cell(0, 7))

        for i in range(len(sites)):

            site = sites[i]
            t.cell(i * 2 + 2, 0).text = site

            ds = dfp.loc[site]

            t.cell(i * 2 + 2, 1).text = 'GPP'
            t.cell(i * 2 + 3, 1).text = 'ET'

            for j in range(2):

                for k in range(2):

                    if k == 0:
                        bloc = 2

                    else:
                        bloc = 5

                    t.cell(i * 2 + j + 2, bloc).text = \
                        '{0:.2f}'.format(score[i, 0, j, k] + score[i, 1, j, k])

                    if conf == 'optimal':
                        t.cell(i * 2 + j + 2, bloc + 1).text = \
                            '{0:.2f}'.format(score[i, len(score[0])-1, j, k])

                        if kmax_calib is not None:
                            t.cell(i * 2 + j + 2, bloc + 2).text = \
                                '{0:.2f}'.format(score[i, 2, j, k])

        if conf == 'optimal':
            t.cell(len(sites) * 2 + 2, 0).text = 'across sites mean'
            t.cell(len(sites) * 2 + 2, 1).text = 'GPP'
            t.cell(len(sites) * 2 + 3, 1).text = 'ET'

            for j in range(2):  # GPP, ET

                for k in range(2):  # events

                    if k == 0:
                        bloc = 2

                    else:
                        bloc = 5

                    t.cell(len(sites) * 2 + j + 2, bloc).text = \
                        '{0:.2f}'.format(np.mean(score[:, 0, j, k] +
                                                 score[:, 1, j, k]))
                    t.cell(len(sites) * 2 + j + 2, bloc + 1).text = \
                        ('{0:.2f}'
                         .format(np.mean(score[:, len(score[0])-1, j, k])))

                    if kmax_calib is not None:
                        t.cell(len(sites) * 2 + j + 2, bloc + 2).text = \
                            '{0:.2f}'.format(np.mean(score[:, 2, j, k]))

            t.cell((len(sites) + 1) * 2 + 2, 0).text = 'across sites median'
            t.cell((len(sites) + 1) * 2 + 2, 1).text = 'GPP'
            t.cell((len(sites) + 1) * 2 + 3, 1).text = 'ET'

            for j in range(2):  # GPP, ET

                for k in range(2):  # events

                    if k == 0:
                        bloc = 2

                    else:
                        bloc = 5

                    t.cell((len(sites) + 1) * 2 + j + 2, bloc).text = \
                        '{0:.2f}'.format(np.median(score[:, 0, j, k] +
                                                   score[:, 1, j, k]))
                    t.cell((len(sites) + 1) * 2 + j + 2, bloc + 1).text = \
                        '{0:.2f}'.format(np.median(score[:,
                                                   len(score[0])-1, j, k]))

                    if kmax_calib is not None:
                        t.cell((len(sites) + 1) * 2 + j + 2, bloc + 2).text = \
                            '{0:.2f}'.format(np.median(score[:, 2, j, k]))

            t.cell((len(sites) + 2) * 2 + 3, 0).text = 'overall mean'
            t.cell((len(sites) + 2) * 2 + 3, 1).text = 'GPP'
            t.cell((len(sites) + 2) * 2 + 4, 1).text = 'ET'

            for j in range(2):  # GPP, ET

                t.cell((len(sites) + 2) * 2 + j + 3, 2).text = \
                    '{0:.2f}'.format(np.mean(score[:, 0, j, :] +
                                             score[:, 1, j, :]))
                t.cell((len(sites) + 2) * 2 + j + 3, 3).text = \
                    '{0:.2f}'.format(np.mean(score[:, len(score[0])-1, j, :]))

                if kmax_calib is not None:
                    t.cell((len(sites) + 2) * 2 + j + 3, 4).text = \
                        '{0:.2f}'.format(np.mean(score[:, 2, j, :]))

            t.cell((len(sites) + 3) * 2 + 3, 0).text = 'overall median'
            t.cell((len(sites) + 3) * 2 + 3, 1).text = 'GPP'
            t.cell((len(sites) + 3) * 2 + 4, 1).text = 'ET'

            for j in range(2):  # GPP, ET

                t.cell((len(sites) + 3) * 2 + j + 3, 2).text = \
                    '{0:.2f}'.format(np.median(score[:, 0, j, :] +
                                               score[:, 1, j, :]))
                t.cell((len(sites) + 3) * 2 + j + 3, 3).text = \
                    '{0:.2f}'.format(np.median(score[:,
                                                     len(score[0])-1, j, :]))

                if kmax_calib is not None:
                    t.cell((len(sites) + 3) * 2 + j + 3, 4).text = \
                        '{0:.2f}'.format(np.median(score[:, 2, j, :]))

        track += 1

    # retrieve the total values
    GPP_avg, ET_avg, ratio1_ctrl, ratio2_ctrl = [None, None], [None, None], \
                                                [None, None], [None, None]
    GPP_xx, ET_xx = [None, None], [None, None]
    GPP_cal, ET_cal, ratio1_cal, ratio2_cal = [None, None], [None, None], \
                                              [None, None], [None, None]
    GPP, ET = [None, None], [None, None]

    track = 0

    for y in [[2003, 2006], [2002, 2005]]:

        GPP[track], ET[track] = fluxes(sites, 'var_kmax', years=y)

        for project in projects:

            if 'average' in project:
                GPP_avg[track], ET_avg[track], ratio1_ctrl[track], \
                    ratio2_ctrl[track] = fluxes(sites, project, years=y,
                                                which='sim')

            if 'extreme' in project:
                GPP_xx[track], ET_xx[track], __, __ = fluxes(sites, project,
                                                             years=y,
                                                             which='sim')

            if 'sample' in project:
                ksites = [sites[i] + e for i, e in enumerate(kmax_calib)]
                GPP_cal[track], ET_cal[track], ratio1_cal[track], \
                    ratio2_cal[track] = fluxes(ksites, project, years=y,
                                               which='calib')

        track += 1

    doc.add_heading('Total GPP April-November (g C/m2)', level=1)

    t = doc.add_table(rows=len(sites) * 3 + 1, cols=7)
    t.cell(0, 0).text = 'Site Name'
    t.cell(0, 1).text = 'years'
    t.cell(0, 2).text = 'average scenario'
    t.cell(0, 3).text = 'extreme scenario'
    t.cell(0, 4).text = 'control'
    t.cell(0, 5).text = 'calib'
    t.cell(0, 6).text = 'obs'

    for i in range(len(sites)):

        site = sites[i]
        t.cell(i * 3 + 1, 0).text = site

        track = 0

        for y in [[2003, 2006], [2002, 2005], 'across years']:

            if y == [2003, 2006]:
                loc = 1

            elif y == [2002, 2005]:
                loc = 2

            else:
                loc = 3

            if y != 'across years':
                t.cell(i * 3 + loc, 1).text = '%d & %d' % (y[0], y[1])
                t.cell(i * 3 + loc, 2).text = \
                    '{0:.2f}'.format(GPP_avg[track][i][1][1])
                t.cell(i * 3 + loc, 3).text = \
                    '{0:.2f}'.format(GPP_xx[track][i][1][1])
                t.cell(i * 3 + loc, 4).text = \
                    '{0:.2f}'.format(GPP_avg[track][i][1][0])
                t.cell(i * 3 + loc, 5).text = \
                    '{0:.2f}'.format(GPP_cal[track][i])
                t.cell(i * 3 + loc, 6).text = '{0:.2f}'.format(GPP[track][i])

            else:
                t.cell(i * 3 + loc, 1).text = 'across years'
                t.cell(i * 3 + loc, 2).text = \
                    '{0:.2f}'.format((GPP_avg[0][i][1][1] +
                                      GPP_avg[1][i][1][1]) / 2.)
                t.cell(i * 3 + loc, 3).text = \
                    '{0:.2f}'.format((GPP_xx[0][i][1][1] +
                                      GPP_xx[1][i][1][1]) / 2.)
                t.cell(i * 3 + loc, 4).text = \
                    '{0:.2f}'.format((GPP_avg[0][i][1][0] +
                                      GPP_avg[1][i][1][0]) / 2.)
                t.cell(i * 3 + loc, 5).text = \
                    '{0:.2f}'.format((GPP_cal[0][i] + GPP_cal[1][i]) / 2.)
                t.cell(i * 3 + loc, 6).text = \
                    '{0:.2f}'.format((GPP[0][i] + GPP[1][i]) / 2.)

            track += 1

    doc.add_heading('% total April-Nov improvement on GPP Control', level=1)

    t = doc.add_table(rows=len(sites) + 1, cols=4)
    t.cell(0, 0).text = 'Site Name'
    t.cell(0, 1).text = 'average scenario'
    t.cell(0, 2).text = 'extreme scenario'
    t.cell(0, 3).text = 'calib'

    for i in range(len(sites)):

        site = sites[i]
        t.cell(i + 1, 0).text = site

        obs = (GPP[0][i] + GPP[1][i]) / 2.
        ctrl = (GPP_avg[0][i][1][0] + GPP_avg[1][i][1][0]) / 2.
        avg = (GPP_avg[0][i][1][1] + GPP_avg[1][i][1][1]) / 2.
        xx = (GPP_xx[0][i][1][1] + GPP_xx[1][i][1][1]) / 2.
        cal = (GPP_cal[0][i] + GPP_cal[1][i]) / 2.

        avg = (abs(obs - ctrl) - abs(obs - avg)) / abs(obs - ctrl)
        xx = (abs(obs - ctrl) - abs(obs - xx)) / abs(obs - ctrl)
        cal = (abs(obs - ctrl) - abs(obs - cal)) / abs(obs - ctrl)

        t.cell(i + 1, 1).text = '{0:.2f}'.format(100. * avg)
        t.cell(i + 1, 2).text = '{0:.2f}'.format(100. * xx)
        t.cell(i + 1, 3).text = '{0:.2f}'.format(100. * cal)

    doc.add_heading('Total ET April-November (mm)', level=1)

    t = doc.add_table(rows=len(sites) * 3 + 1, cols=7)
    t.cell(0, 0).text = 'Site Name'
    t.cell(0, 1).text = 'years'
    t.cell(0, 2).text = 'average scenario'
    t.cell(0, 3).text = 'extreme scenario'
    t.cell(0, 4).text = 'control'
    t.cell(0, 5).text = 'calib'
    t.cell(0, 6).text = 'obs'

    for i in range(len(sites)):

        site = sites[i]
        t.cell(i * 3 + 1, 0).text = site

        track = 0

        for y in [[2003, 2006], [2002, 2005], 'across years']:

            if y == [2003, 2006]:
                loc = 1

            elif y == [2002, 2005]:
                loc = 2

            else:
                loc = 3

            if y != 'across years':
                t.cell(i * 3 + loc, 1).text = '%d & %d' % (y[0], y[1])
                t.cell(i * 3 + loc, 2).text = \
                    '{0:.2f}'.format(ET_avg[track][i][1][1])
                t.cell(i * 3 + loc, 3).text = \
                    '{0:.2f}'.format(ET_xx[track][i][1][1])
                t.cell(i * 3 + loc, 4).text = \
                    '{0:.2f}'.format(ET_avg[track][i][1][0])
                t.cell(i * 3 + loc, 5).text = \
                    '{0:.2f}'.format(ET_cal[track][i])
                t.cell(i * 3 + loc, 6).text = '{0:.2f}'.format(ET[track][i])

            else:
                t.cell(i * 3 + loc, 1).text = 'across years'
                t.cell(i * 3 + loc, 2).text = \
                    '{0:.2f}'.format((ET_avg[0][i][1][1] +
                                      ET_avg[1][i][1][1]) / 2.)
                t.cell(i * 3 + loc, 3).text = \
                    '{0:.2f}'.format((ET_xx[0][i][1][1] +
                                      ET_xx[1][i][1][1]) / 2.)
                t.cell(i * 3 + loc, 4).text = \
                    '{0:.2f}'.format((ET_avg[0][i][1][0] +
                                      ET_avg[1][i][1][0]) / 2.)
                t.cell(i * 3 + loc, 5).text = \
                    '{0:.2f}'.format((ET_cal[0][i] + ET_cal[1][i]) / 2.)
                t.cell(i * 3 + loc, 6).text = \
                    '{0:.2f}'.format((ET[0][i] + ET[1][i]) / 2.)

            track += 1

    doc.add_heading('% total April-Nov improvement on ET Control', level=1)

    t = doc.add_table(rows=len(sites) + 1, cols=4)
    t.cell(0, 0).text = 'Site Name'
    t.cell(0, 1).text = 'average scenario'
    t.cell(0, 2).text = 'extreme scenario'
    t.cell(0, 3).text = 'calib'

    for i in range(len(sites)):

        site = sites[i]
        t.cell(i + 1, 0).text = site

        obs = (ET[0][i] + ET[1][i]) / 2.
        ctrl = (ET_avg[0][i][1][0] + ET_avg[1][i][1][0]) / 2.
        avg = (ET_avg[0][i][1][1] + ET_avg[1][i][1][1]) / 2.
        xx = (ET_xx[0][i][1][1] + ET_xx[1][i][1][1]) / 2.
        cal = (ET_cal[0][i] + ET_cal[1][i]) / 2.

        avg = (abs(obs - ctrl) - abs(obs - avg)) / abs(obs - ctrl)
        xx = (abs(obs - ctrl) - abs(obs - xx)) / abs(obs - ctrl)
        cal = (abs(obs - ctrl) - abs(obs - cal)) / abs(obs - ctrl)

        t.cell(i + 1, 1).text = '{0:.2f}'.format(100. * avg)
        t.cell(i + 1, 2).text = '{0:.2f}'.format(100. * xx)
        t.cell(i + 1, 3).text = '{0:.2f}'.format(100. * cal)

    doc.add_heading('Es:ET & E:ET average instantaneous ratios April-November',
                    level=1)

    t = doc.add_table(rows=len(sites) * 3 + 1, cols=6)
    t.cell(0, 0).text = 'Site Name'
    t.cell(0, 1).text = 'years'
    t.cell(0, 2).text = 'control Es:ET'
    t.cell(0, 3).text = 'calib Es:ET'
    t.cell(0, 4).text = 'control E:ET'
    t.cell(0, 5).text = 'calib E:ET'

    for i in range(len(sites)):

        site = sites[i]
        t.cell(i * 3 + 1, 0).text = site

        track = 0

        for y in [[2003, 2006], [2002, 2005], 'across years']:

            if y == [2003, 2006]:
                loc = 1

            elif y == [2002, 2005]:
                loc = 2

            else:
                loc = 3

            if y != 'across years':
                t.cell(i * 3 + loc, 1).text = '%d & %d' % (y[0], y[1])
                t.cell(i * 3 + loc, 2).text = \
                    '{0:.2f}'.format(ratio1_ctrl[track][i][1][0])
                t.cell(i * 3 + loc, 3).text = \
                    '{0:.2f}'.format(ratio1_cal[track][i])
                t.cell(i * 3 + loc, 4).text = \
                    '{0:.2f}'.format(ratio2_ctrl[track][i][1][0])
                t.cell(i * 3 + loc, 5).text = \
                    '{0:.2f}'.format(ratio2_cal[track][i])

            else:
                t.cell(i * 3 + loc, 1).text = 'across years'
                t.cell(i * 3 + loc, 2).text = \
                    '{0:.2f}'.format((ratio1_ctrl[0][i][1][0] +
                                      ratio1_ctrl[1][i][1][0]) / 2.)
                t.cell(i * 3 + loc, 3).text = \
                    '{0:.2f}'.format((ratio1_cal[0][i] +
                                      ratio1_cal[1][i]) / 2.)
                t.cell(i * 3 + loc, 4).text = \
                    '{0:.2f}'.format((ratio2_ctrl[0][i][1][0] +
                                      ratio2_ctrl[1][i][1][0]) / 2.)
                t.cell(i * 3 + loc, 5).text = \
                    '{0:.2f}'.format((ratio2_cal[0][i] +
                                      ratio2_cal[1][i]) / 2.)

            track += 1

    doc.save(fname)

    return


#=======================================================================

if __name__ == "__main__":

    # define the argparse settings to read run set up file
    description = "Plot all the different site fluxes & soil variations for a \
                   given project folder"
    parser = argparse.ArgumentParser(description=description)
    parser.add_argument("project", nargs='+', help='projects where output are')
    args = parser.parse_args()

    projects = []

    for arg in vars(args):

        for attr in getattr(args, arg):

            projects += [attr]

    sites = ['Hyytiala', 'Soroe', 'Loobos', 'Hesse', 'Parco', 'Puechabon',
             'Rocca1', 'Rocca2', 'ElSaler1', 'Espirra']

    main(projects)
