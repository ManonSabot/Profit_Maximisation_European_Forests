#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Writes a word document which contains information about outputs and performance
of the calibrated Control model.

This file is part of the TractLSM project.

Copyright (c) 2019 Manon E. B. Sabot

Please refer to the terms of the MIT License, which you should have
received along with the TractLSM.

"""

__title__ = "summary document of the model's configuration and outputs"
__author__ = "Manon E. B. Sabot"
__version__ = "1.0 (15.11.2019)"
__email__ = "m.e.b.sabot@gmail.com"

#=======================================================================

import warnings  # ignore these warnings
warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=RuntimeWarning)

# import general modules
import argparse  # read in the user input
import os  # check for files, paths
import sys  # check for files, paths
import numpy as np  # array manipulations, math operators
import numpy.ma as ma  # masked arrays
import pandas as pd  # read/write dataframes, csv files
import docx  # write out a word document
from docx import Document  # write out a word document

# first make sure that modules can be loaded from TractLSM
script_dir = os.path.dirname(os.path.realpath(sys.argv[0]))
sys.path.append(os.path.abspath(os.path.join(script_dir, '..')))

# own modules
from plotting_scripts.plot_utils import find_site_data  # locate
from TractLSM import conv  # unit converter
from TractLSM.Utils import get_main_dir, read_csv  # locate & read


#=======================================================================

def main(projects):

    """
    Main function: writes a word document with information on general
                   analysis about the output generated.

    Arguments:
    ----------
    projects: array
        names of the repository which contain the output to analyse

    Returns:
    --------
    'summary_calibrating_control.docx' in the output directory

    """

    write_word_doc(os.path.join('output', 'summary_calibrating_control.docx'),
                   projects)

    return


#=======================================================================

# ~~~ Other functions are defined here ~~~

def get_best_param(sites, basedir):

    """
    Retrieves the best site-calibrated parameters.

    Arguments:
    ----------
    sites: array
        names of the sites

    basedir: string
        name (path) of the base repository upon which the variation of
        the parameter was performed

    Returns:
    --------
    params: array
        number names associated with the best performing config for each
        site

    """

    try:  # look for "best" param
        fp = open(os.path.join(basedir, 'best.txt'), 'r')
        flines = fp.readlines()
        param = [None] * len(sites)

        for l in flines:

            site = l.split(':')[0].strip()

            if site in sites:
                p = l.split(':')[1].strip()
                param[sites.index(site)] = p

        if not any(param):
            param = None

    except IOError:
        param = None

    return param


def calculate_scores(data):

    """
    Retrieves each of the yearly metrics for all the sites and fluxes and
    projects.

    Arguments:
    ----------
    data: array
        contains all the paths to the input and output data for which these
        need to be compiled

    Returns:
    --------
    NMSE: array
        normalised mean square errors for each of the variables

    MAE: array
        mean absolute error

    SD: array
        standard deviation

    P5: array
        5th percentile

    P95: array
        95th percentile

    """

    # empty arrays to store the performance scores in
    # dims: sites, projects, fluxes, years
    NMSE = np.zeros((data.shape[2], data.shape[1], 2, 2))
    MAE = np.zeros((data.shape[2], data.shape[1], 2, 2))
    SD = np.zeros((data.shape[2], data.shape[1], 2, 2))
    P5 = np.zeros((data.shape[2], data.shape[1], 2, 2))
    P95 = np.zeros((data.shape[2], data.shape[1], 2, 2))

    # loop over sites
    for j in range(data.shape[2]):

        # loop over projects
        for k in range(data.shape[1]):

            # loop over years
            for l in range(data.shape[3]):

                fluxes = ['GPP', 'Qle']

                try:
                    obs, __ = read_csv(data[0, k, j, l])
                    sim, __ = read_csv(data[1, k, j, l])
                    obs.fillna(0., inplace=True)
                    sim.fillna(0., inplace=True)

                    # on a half-hourly basis
                    Dsteps = int(24. / (obs.loc[1, 'hod'] - obs.loc[0, 'hod']))

                    # restrict to subsets of dates
                    year = int(data[0, k, j, l].split('_')[-1]
                               .split('.csv')[0])
                    start = pd.to_datetime(year * 1000 + obs['doy'].iloc[0],
                                           format='%Y%j')

                    if obs['doy'].iloc[-1] > obs['doy'].iloc[0]:
                        end = pd.to_datetime(year * 1000 + obs['doy'].iloc[-1],
                                             format='%Y%j')

                    else:
                        end = pd.to_datetime((year + 1) * 1000 +
                                             obs['doy'].iloc[-1],
                                             format='%Y%j')

                    dates = pd.date_range(start=start, end=end, freq='D')
                    __, iunique = np.unique(dates.month, return_index=True)
                    unique_months = [dates.month[m] for m in sorted(iunique)]

                    # all, restrict dates to growing season
                    dates_2_keep = np.in1d(dates.month, unique_months[3:-2])

                    # restrict to day time and acceptable LAI
                    mask = np.logical_and(obs['PPFD'] > 50.,
                                          obs['LAI'] > 0.001)
                    obs = obs.where(mask).fillna(value=0.)
                    sim = sim.where(mask).fillna(value=0.)

                    # unit conversion
                    obs['Qle'] *= conv.Wpm2_2_mmphlfhr
                    sim['Qle'] = ((sim['E(std)'] + sim['Es(std)']) *
                                  conv.mmolH2Opm2ps_2_mmphlfhr)
                    obs['GPP'] *= conv.umolCpm2ps_2_gCpm2phlfhr
                    sim['GPP'] = sim['A(std)'] * conv.umolCpm2ps_2_gCpm2phlfhr

                    # daily sums
                    obs = obs[fluxes]
                    sim = sim[fluxes]
                    obs = obs.groupby(obs.index // Dsteps * Dsteps).sum()
                    sim = sim.groupby(sim.index // Dsteps * Dsteps).sum()
                    obs.reset_index(inplace=True)
                    sim.reset_index(inplace=True)
                    obs.drop(columns=['index'], inplace=True)
                    sim.drop(columns=['index'], inplace=True)

                    # restricted dates
                    obs = obs[dates_2_keep]
                    sim = sim[dates_2_keep]

                    # deal with missing or negative values
                    mask = np.logical_and(obs > 0., sim > 0.)
                    obs = obs[mask]
                    sim = sim[mask]
                    obs.fillna(0., inplace=True)
                    sim.fillna(0., inplace=True)

                    # loop over An and ET
                    for i in range(len(fluxes)):

                        subobs = obs[fluxes[i]]
                        subsim = sim[fluxes[i]]

                        nmse = np.mean((subsim - subobs) ** 2. /
                                       (np.mean(subsim) * np.mean(subobs)))
                        mae = np.mean(np.abs(subsim - subobs))
                        sd = np.abs(1. - subsim.std() /
                                    np.maximum(1.e-9, subobs.std()))
                        p5 = np.abs(np.percentile(subsim, 5) -
                                    np.percentile(subobs, 5))
                        p95 = np.abs(np.percentile(subsim, 95) -
                                     np.percentile(subobs, 95))

                        if l > 1:
                            NMSE[j, k, i, l-2] = 0.5 * (NMSE[j, k, i, l-2] +
                                                        nmse)
                            MAE[j, k, i, l-2] = 0.5 * (MAE[j, k, i, l-2] + mae)
                            SD[j, k, i, l-2] = 0.5 * (SD[j, k, i, l-2] + sd)
                            P5[j, k, i, l-2] = 0.5 * (P5[j, k, i, l-2] + p5)
                            P95[j, k, i, l-2] = 0.5 * (P95[j, k, i, l-2] + p95)

                        else:
                            NMSE[j, k, i, l] = nmse
                            MAE[j, k, i, l] = mae
                            SD[j, k, i, l] = sd
                            P5[j, k, i, l] = p5
                            P95[j, k, i, l] = p95

                except IOError:
                    pass

    return NMSE, MAE, SD, P5, P95


def write_word_doc(fname, projects):

    """
    Writes the final word document. The order of the tables is as
    follows:

        1. Control calibration NMSEs (close to zero is good)
        2. Control calibration MAEs (close to zero is good)
        3. Control calibration SD metrics (close to zero is good)
        4. Control calibration P5 metrics (close to zero is good)
        5. Control calibration P95 metrics (close to zero is good)
        6. % improvement on the Calibration metrics compared to the
            Control

    Arguments:
    ----------
    fname: string
        name of the summary file to save

    projects: array
        names of the repository which contain the kmax data

    Returns:
    --------
    fname in the directory which harbours TractLSM

    """

    # useful directories
    basedir = get_main_dir()

    while 'src' in basedir:
        basedir = os.path.dirname(get_main_dir())

    Rout = os.path.join(os.path.join(basedir, 'output'), 'projects')

    # declare document
    doc = Document()

    doc.add_heading("Summary of the calibration of the Control", 0)

    description = """This document summarizes:\n
                        general analysis about the output generated
                  """
    p = doc.add_paragraph(description)
    p.add_run('\nThe following projects are accounted for:\n')

    # all file names, dims: idata/odata, projects, sites, years
    data = np.empty((2, len(projects), len(sites), 4), dtype='|S1000')

    tr1 = 0

    for project in projects:

        p.add_run('     %s\n' % (project)).bold = True
        cal = get_best_param(sites, os.path.join(Rout, project))

        if 'g1' in project:
            g1 = cal

        elif 'fw' in project:
            fw = cal

        else:
            cal = ['', ] * len(sites)

        # find all related data
        for i in range(len(sites)):

            tr2 = 0

            for year in [2002, 2003, 2005, 2006]:

                idata, odata = find_site_data(sites[i] + cal[i], year, project)

                try:
                    data[0, tr1, i, tr2] = "%s" % (idata[0])
                    data[1, tr1, i, tr2] = "%s" % (odata[0])

                except Exception:
                    pass

                tr2 += 1

        tr1 += 1

    # performance scores
    scores = (calculate_scores(data))

    track = 0

    for perf in ['NMSE', 'MAE', 'SD', 'P5', 'P95']:

        doc.add_heading('Control optimal %ss' % (perf), level=1)
        score = scores[track]

        t = doc.add_table(rows=len(sites) * 2 + 7, cols=8)
        t.cell(0, 0).text = 'Site Name'
        t.cell(0, 1).text = 'Flux'
        t.cell(0, 2).text = '2002 & 2005'
        t.cell(0, 5).text = '2003 & 2006'
        t.cell(1, 2).text = 'calib g1'
        t.cell(1, 3).text = 'calib fw'
        t.cell(1, 4).text = 'control'
        t.cell(1, 5).text = 'calib g1'
        t.cell(1, 6).text = 'calib fw'
        t.cell(1, 7).text = 'control'

        # rendering of the headers
        t.cell(0, 0).merge(t.cell(1, 0))
        t.cell(0, 2).merge(t.cell(0, 3))
        t.cell(0, 3).merge(t.cell(0, 4))
        t.cell(0, 5).merge(t.cell(0, 6))
        t.cell(0, 6).merge(t.cell(0, 7))

        for i in range(len(sites)):

            site = sites[i]
            t.cell(i * 2 + 2, 0).text = site
            t.cell(i * 2 + 2, 1).text = 'GPP'
            t.cell(i * 2 + 3, 1).text = 'ET'

            for j in range(2):  # GPP, ET

                for k in range(2):  # events

                    if k == 0:
                        bloc = 2

                    else:
                        bloc = 5

                    t.cell(i * 2 + j + 2, bloc).text = \
                        '{0:.2f}'.format(score[i, 0, j, k])
                    t.cell(i * 2 + j + 2, bloc + 1).text = \
                        '{0:.2f}'.format(score[i, 1, j, k])
                    t.cell(i * 2 + j + 2, bloc + 2).text = \
                        '{0:.2f}'.format(score[i, 2, j, k])

        t.cell(len(sites) * 2 + 2, 0).text = 'across sites'
        t.cell(len(sites) * 2 + 2, 1).text = 'GPP'
        t.cell(len(sites) * 2 + 3, 1).text = 'ET'

        for j in range(2):  # GPP, ET

            for k in range(2):  # events

                if k == 0:
                    bloc = 2

                else:
                    bloc = 5

                t.cell(len(sites) * 2 + j + 2, bloc).text = \
                    '{0:.2f}'.format(np.mean(score[:, 0, j, k]))
                t.cell(len(sites) * 2 + j + 2, bloc + 1).text = \
                    '{0:.2f}'.format(np.mean(score[:, 1, j, k]))
                t.cell(len(sites) * 2 + j + 2, bloc + 2).text = \
                    '{0:.2f}'.format(np.mean(score[:, 2, j, k]))

        t.cell((len(sites) + 1) * 2 + 3, 0).text = 'overall'
        t.cell((len(sites) + 1) * 2 + 3, 1).text = 'GPP'
        t.cell((len(sites) + 1) * 2 + 4, 1).text = 'ET'

        for j in range(2):  # GPP, ET

            t.cell((len(sites) + 1) * 2 + j + 3, 2).text = \
                '{0:.2f}'.format(np.mean(score[:, 0, j, :]))
            t.cell((len(sites) + 1) * 2 + j + 3, 3).text = \
                '{0:.2f}'.format(np.mean(score[:, 1, j, :]))
            t.cell((len(sites) + 1) * 2 + j + 3, 4).text = \
                '{0:.2f}'.format(np.mean(score[:, 2, j, :]))

        track += 1

    doc.add_heading('% overall improvement on Control metrics', level=1)

    t = doc.add_table(rows=11, cols=4)
    t.cell(0, 0).text = 'metric'
    t.cell(0, 1).text = 'flux'
    t.cell(0, 2).text = 'g1'
    t.cell(0, 3).text = 'fw'

    track1 = 0
    track2 = 1

    for perf in ['NMSE', 'MAE', 'SD', 'P5', 'P95']:

        score = ma.filled(scores[track1], 0.)  # 0. where masked values

        t.cell(track2, 0).text = perf
        t.cell(track2, 0).merge(t.cell(track2 + 1, 0))

        for j in range(2):  # GPP, ET

            if j == 0:
                t.cell(track2, 1).text = 'GPP'

            else:
                t.cell(track2, 1).text = 'ET'

            ctrl = np.mean(score[:, 2, j, :])
            g1 = np.mean(score[:, 0, j, :])
            g1 = (abs(ctrl) - abs(g1)) / abs(ctrl)
            t.cell(track2, 2).text = '{0:.2f}'.format(100. * g1)
            fw = np.mean(score[:, 1, j, :])
            fw = (abs(ctrl) - abs(fw)) / abs(ctrl)
            t.cell(track2, 3).text = '{0:.2f}'.format(100. * fw)

            track2 += 1

        track1 += 1

    doc.save(fname)

    return


#=======================================================================

if __name__ == "__main__":

    # define the argparse settings to read run set up file
    description = "Plot all the different site fluxes & soil variations for a \
                   given project folder"
    parser = argparse.ArgumentParser(description=description)
    parser.add_argument("project", nargs='+', help='projects where output are')
    args = parser.parse_args()

    projects = []

    for arg in vars(args):

        for attr in getattr(args, arg):

            projects += [attr]

    sites = ['Hyytiala', 'Soroe', 'Loobos', 'Hesse', 'Parco', 'Puechabon',
             'Rocca1', 'Rocca2', 'ElSaler1', 'Espirra']

    main(projects)
